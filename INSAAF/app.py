from flask import Flask, render_template, request, send_file, redirect, url_for
from werkzeug.utils import secure_filename
import os
from docx import Document
from hashlib import sha256

app = Flask(__name__)

# Configuration for file uploads
UPLOAD_FOLDER = 'uploads/'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# In-memory storage for user credentials and files (for demo purposes)
users = {}

# Function to create the document
def create_document(form_data):
    doc = Document()

    doc.add_heading('Criminal Miscellaneous Writ Petition', 0)

    doc.add_paragraph(f"Petitioner: {form_data['petitionerName']}")
    doc.add_paragraph(f"Respondent: {form_data['respondentName']}")
    doc.add_paragraph(f"Case Details:\n{form_data['caseDetails']}")

    filename = os.path.join(app.config['UPLOAD_FOLDER'], f"{form_data['petitionerName']}_petition.docx")
    doc.save(filename)
    return filename

@app.route('/submit_plaint', methods=['POST'])
def submit_plaint():
    # Get form data
    petitioner_name = request.form['petitionerName']
    respondent_name = request.form['respondentName']
    case_details = request.form['caseDetails']
    password = request.form['password']

    # Secure password
    hashed_password = sha256(password.encode()).hexdigest()

    # Create the document
    filename = create_document({
        'petitionerName': petitioner_name,
        'respondentName': respondent_name,
        'caseDetails': case_details
    })

    # Save identity proof and supporting documents
    identity_proof = request.files['identityProof']
    identity_proof_filename = secure_filename(identity_proof.filename)
    identity_proof.save(os.path.join(app.config['UPLOAD_FOLDER'], identity_proof_filename))

    supporting_docs = request.files.getlist('supportingDocuments')
    supporting_docs_filenames = []
    for doc in supporting_docs:
        doc_filename = secure_filename(doc.filename)
        doc.save(os.path.join(app.config['UPLOAD_FOLDER'], doc_filename))
        supporting_docs_filenames.append(doc_filename)

    # Store user info
    users[petitioner_name] = {
        'password': hashed_password,
        'petition_file': filename,
        'identity_proof': identity_proof_filename,
        'supporting_docs': supporting_docs_filenames
    }

    return send_file(filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)